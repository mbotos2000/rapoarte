import streamlit as st
import ftplib
import pickle
import io
import pandas as pd
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, nsmap

def set_table_autofit(table):
    """Ensures the table fits its content."""
    tbl = table._element
    tblPr = tbl.find(".//w:tblPr", nsmap)
    
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblLayout = tblPr.find(".//w:tblLayout", nsmap)
    
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)

    tblLayout.set("w:type", "autofit")  # Ensures the table resizes dynamically

def generate_docx_with_table(dataframe, titlu):
    """Generates a DOCX file with a table from a DataFrame and returns it as a BytesIO object."""
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # Add title
    doc.add_heading(titlu, level=1)
    
    # Create table with the same number of columns as the DataFrame
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"  # You can change the style
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column)

    # Add data rows
    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    
    # Apply autofit to make the table fit its content
    set_table_autofit(table)

    # Save the DOCX file to an in-memory bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
@st.cache_data(show_spinner=False)
def load_data_from_ftp():
    
    ftp = ftplib.FTP("users.utcluj.ro", st.secrets['u'], st.secrets['p'])
    ftp.encoding = "utf-8"  # Force UTF-8 encoding
    ftp.cwd('/public_html/Fise')

    # List all files in the directory
    files = ftp.nlst()

    # Initialize a list to store all dictionaries
    all_dicts = []

    # Loop through each file
    for file_name in files:
        if file_name.endswith('.pkl'):
            # Create a BytesIO object to hold the file data
            file_data = io.BytesIO()
            
            # Download the file from the FTP server
            ftp.retrbinary(f'RETR {file_name}', file_data.write)
            
            # Seek to the beginning of the BytesIO object
            file_data.seek(0)
            
            # Load the dictionary from the .pkl file
            dictionary = pickle.load(file_data)
            
            # Append the dictionary to the list
            all_dicts.append(dictionary)
            
            # Close the BytesIO object
            file_data.close()

    # Close the FTP connection
    ftp.quit()

    # Combine all dictionaries into a single DataFrame
    d = pd.DataFrame(all_dicts)
    return d

def main():
    st.title("Generator rapoarte specializari")
    
    df = load_data_from_ftp()

    if df is not None:
        app_col=['M_8_2_1','M_8_2_2','M_8_2_3','M_8_2_4','M_8_2_5','M_8_2_6','M_8_2_7',
                 'M_8_2_8','M_8_2_9','M_8_2_10','M_8_2_11','M_8_2_12','M_8_2_13','M_8_2_14']
        curs_col=['M_8_1_1','M_8_1_2','M_8_1_3','M_8_1_4','M_8_1_5','M_8_1_6','M_8_1_7',
                 'M_8_1_8','M_8_1_9','M_8_1_10','M_8_1_11','M_8_1_12','M_8_1_13','M_8_1_14']
        df['Cursuri']=df[curs_col].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Aplicatii']=df[app_col].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Preconditii']=df[['M_4_1','M_4_2']].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Conditii']=df[['M_5_1','M_5_2']].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Competente']=df[['M_6_cp','M_6_ct']].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Titulari']=df[['M_2_2','M_2_3']].agg(lambda x: ' '.join(x) + '\n', axis=1)
        df['Obiective']=df[['M_7_1','M_7_2']].agg(lambda x: ' '.join(x) + '\n', axis=1)
        unique_values = df['M_1_6'].dropna().unique()
        specializarea = st.selectbox("Selecteaza scpecializarea ", list(unique_values))
        if specializarea!=None:
            filtered_df = df.copy()
            filtered_df = filtered_df.sort_values(by='M_1_8')
            filtered_df = filtered_df[filtered_df['M_1_6'] == specializarea]
            categories = df["M_2_7_1"].unique().tolist()
            selected_categories = st.multiselect("Filtreaza dupa tipul disciplinei", categories, default=categories)
            categories_1 = df["M_2_7_2"].unique().tolist()
            selected_categories_1 = st.multiselect("Filtreaza dupa regimul disciplinei", categories_1, default=categories_1)
            categories_2 = df["M_2_4"].unique().tolist()
            selected_categories_2 = st.multiselect("Filtreaza dupa anul de studiu", categories_2, default=categories_2)
            if st.button("Aplica filtrele suplimentare"):
                filtered_df = filtered_df[filtered_df["M_2_7_1"].isin(selected_categories)]
                filtered_df = filtered_df[filtered_df["M_2_7_2"].isin(selected_categories_1)]
                filtered_df = filtered_df[filtered_df["M_2_4"].isin(selected_categories_2)]
                filtered_df['ordonare']=filtered_df['M_1_8'].astype(float)
                filtered_df = filtered_df.sort_values(by='ordonare')
                filtered_df['Cod disciplina']=filtered_df['M_1_8']
                filtered_df['Denumire disciplina']=filtered_df['M_2_1']
                filtered_df['Titularul de curs']=filtered_df['M_2_2']
                filtered_df['Titular aplicatii']=filtered_df['M_2_3']
                filtered_df['Tipul de evaluare']=filtered_df['M_2_6']
                filtered_df['Regimul disciplinei']=filtered_df['M_2_7_1']
                filtered_df['Numar credite']=filtered_df['M_3_11']
                
                report_df = filtered_df[['Cod disciplina','Denumire disciplina','Cursuri','Aplicatii']]
                report_df_1 = filtered_df[['Cod disciplina','Denumire disciplina','Competente']]
                report_df_2 = filtered_df[['Cod disciplina','Denumire disciplina','Preconditii']]
                report_df_3 = filtered_df[['Cod disciplina','Denumire disciplina','Conditii']]
                report_df_4 = filtered_df[['Cod disciplina','Denumire disciplina','Obiective']]
                report_df_5 = filtered_df[['Cod disciplina','Denumire disciplina','Titulari']]
                #st.write("### Generated Report")
            
        
                if not report_df.empty:
                    st.write("Datele necesare generarii raportelor au fost citite!")
                    #csv = report_df.to_csv(index=False).encode("utf-8-sig")
                    #csv1 = report_df_1.to_csv(index=False).encode("utf-8-sig")
                    #csv2 = report_df_2.to_csv(index=False).encode("utf-8-sig")
                    #csv3 = report_df_3.to_csv(index=False).encode("utf-8-sig")
                    #csv4 = report_df_4.to_csv(index=False).encode("utf-8-sig")
                    #csv5 = report_df_5.to_csv(index=False).encode("utf-8-sig")
                    
                    docx_file = generate_docx_with_table(report_df, "Raport cursuri si aplicatii")
                    docx_file_1 = generate_docx_with_table(report_df_1, "Raport competente")
                    docx_file_2 = generate_docx_with_table(report_df_2, "Raport preconditii")
                    docx_file_3 = generate_docx_with_table(report_df_3, "Raport conditii")
                    docx_file_4 = generate_docx_with_table(report_df_4, "Raport obiective")
                    docx_file_5 = generate_docx_with_table(report_df_5, "Raport cadre didactice")
                    # Create a download button
                    st.download_button(
                        label="Raport continuturi discipline .docx",
                        data=docx_file,
                        file_name="Raport_continuturi.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label="Raport competente .docx",
                        data=docx_file_1,
                        file_name="Raport_competente.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label="Raport preconditii .docx",
                        data=docx_file_2,
                        file_name="Raport_preconditii.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label="Raport conditii .docx",
                        data=docx_file_3,
                        file_name="Raport_conditii.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label="Raport obiective .docx",
                        data=docx_file_4,
                        file_name="Raport_obiective.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.download_button(
                        label="Raport cadre didactice .docx",
                        data=docx_file_5,
                        file_name="Raport_CD.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()
